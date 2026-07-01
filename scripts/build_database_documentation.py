from __future__ import annotations

import csv
from collections import Counter, defaultdict
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("outputs/database_documentation")
PNG_PATH = OUTPUT_DIR / "database_visual_map.png"
DOCX_PATH = OUTPUT_DIR / "database_table_inventory.docx"

SCHEMA_ORDER = ["reference", "access", "planning", "performance", "review", "workflow"]
SCHEMA_COLORS = {
    "reference": "#DCE9F7",
    "access": "#E6F2E6",
    "planning": "#FFF1CC",
    "performance": "#EDE4F7",
    "review": "#F7E2DF",
    "workflow": "#E7EEF0",
}
SCHEMA_PURPOSE = {
    "reference": "Shared reference data: agencies, services, pillars, action plan content, and plan entities.",
    "access": "Users, roles, and agency-level permissions.",
    "planning": "Plan cycle records, plan shells, draft payloads, and budget proposals.",
    "performance": "Agency plans, goals, services, measures, actuals, risks, and entity-aware measure links.",
    "review": "Reviewer assignments, rubric scores, review feedback, and approval recommendations.",
    "workflow": "Status history and audit trail for plan movement through review.",
}


def read_csv(name: str) -> list[dict[str, str]]:
    with (OUTPUT_DIR / name).open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def font(size: int, bold: bool = False):
    paths = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in paths:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def wrap_text(text: str, max_chars: int) -> list[str]:
    words = str(text).split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        candidate = " ".join(current + [word])
        if len(candidate) <= max_chars:
            current.append(word)
        else:
            if current:
                lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines or [""]


def table_key(row: dict[str, str]) -> str:
    return f"{row['table_schema']}.{row['table_name']}"


def build_visual(tables: list[dict[str, str]], foreign_keys: list[dict[str, str]]) -> None:
    width, height = 2400, 1650
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(42, True)
    subtitle_font = font(22)
    schema_font = font(26, True)
    body_font = font(18)
    small_font = font(16)

    draw.text((80, 60), "Beacon Database Visual Map", fill="#2F2140", font=title_font)
    draw.text(
        (80, 115),
        "Schemas, major table groups, and cross-schema relationships in the current local database.",
        fill="#4B5563",
        font=subtitle_font,
    )

    schema_tables = defaultdict(list)
    for row in tables:
        schema_tables[row["table_schema"]].append(row["table_name"])

    positions = {
        "reference": (80, 210, 720, 690),
        "access": (880, 210, 1520, 550),
        "planning": (1680, 210, 2320, 640),
        "performance": (80, 850, 1120, 1550),
        "review": (1280, 850, 1800, 1280),
        "workflow": (1880, 850, 2320, 1210),
    }

    centers = {}
    for schema in SCHEMA_ORDER:
        x1, y1, x2, y2 = positions[schema]
        centers[schema] = ((x1 + x2) // 2, (y1 + y2) // 2)
        fill = hex_to_rgb(SCHEMA_COLORS[schema])
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline="#9CA3AF", width=3)
        draw.text((x1 + 28, y1 + 22), schema.upper(), fill="#111827", font=schema_font)
        draw.text((x1 + 28, y1 + 58), f"{len(schema_tables[schema])} tables", fill="#374151", font=small_font)
        y = y1 + 98
        for table_name in sorted(schema_tables[schema]):
            lines = wrap_text(table_name, 34 if schema != "performance" else 42)
            if y + (len(lines) * 22) + 18 > y2 - 18:
                draw.text((x1 + 28, y), "...", fill="#374151", font=body_font)
                break
            draw.rounded_rectangle((x1 + 24, y, x2 - 24, y + 30 + (len(lines) - 1) * 20), radius=8, fill="white", outline="#CBD5E1")
            for line_idx, line in enumerate(lines):
                draw.text((x1 + 42, y + 7 + line_idx * 20), line, fill="#111827", font=body_font)
            y += 40 + (len(lines) - 1) * 20

    schema_links = Counter((fk["source_schema"], fk["target_schema"]) for fk in foreign_keys if fk["source_schema"] != fk["target_schema"])
    for (source, target), count in schema_links.items():
        if source not in centers or target not in centers:
            continue
        sx, sy = centers[source]
        tx, ty = centers[target]
        draw.line((sx, sy, tx, ty), fill="#6B7280", width=min(8, 2 + count // 4))
        mx, my = (sx + tx) // 2, (sy + ty) // 2
        draw.ellipse((mx - 24, my - 18, mx + 24, my + 18), fill="#2F2140")
        label = str(count)
        bbox = draw.textbbox((0, 0), label, font=small_font)
        draw.text((mx - (bbox[2] - bbox[0]) / 2, my - 10), label, fill="white", font=small_font)

    legend_x, legend_y = 80, 1570
    draw.text((legend_x, legend_y), "Line labels show number of foreign-key relationships between schemas.", fill="#4B5563", font=small_font)
    image.save(PNG_PATH)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def style_table(table, header_fill: str = "#E8EEF5") -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
        if row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(17, 24, 39)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    style_table(table)


def build_docx(tables: list[dict[str, str]], columns: list[dict[str, str]], primary_keys: list[dict[str, str]], foreign_keys: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in [
        ("Title", 22, RGBColor(47, 33, 64)),
        ("Heading 1", 16, RGBColor(47, 33, 64)),
        ("Heading 2", 13, RGBColor(31, 77, 120)),
        ("Heading 3", 11, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Beacon Database Documentation")
    subtitle = document.add_paragraph("Visual map and table inventory generated from the current local PostgreSQL database.")
    subtitle.paragraph_format.space_after = Pt(12)

    document.add_heading("Database Visual Map", level=1)
    document.add_picture(str(PNG_PATH), width=Inches(6.8))

    document.add_heading("Summary", level=1)
    schema_counts = Counter(row["table_schema"] for row in tables)
    summary_rows = [[schema, str(schema_counts[schema]), SCHEMA_PURPOSE.get(schema, "")] for schema in SCHEMA_ORDER if schema in schema_counts]
    add_table(document, ["Schema", "Tables", "Purpose"], summary_rows, widths=[1.1, 0.7, 4.7])

    pk_map = defaultdict(list)
    for row in primary_keys:
        pk_map[f"{row['table_schema']}.{row['table_name']}"].append(row["column_name"])
    col_counts = Counter(table_key(row) for row in columns)
    references = defaultdict(list)
    referenced_by = defaultdict(list)
    for fk in foreign_keys:
        source = f"{fk['source_schema']}.{fk['source_table']}"
        target = f"{fk['target_schema']}.{fk['target_table']}"
        references[source].append(target)
        referenced_by[target].append(source)

    document.add_heading("Table Inventory", level=1)
    inventory_rows = []
    for row in tables:
        key = table_key(row)
        inventory_rows.append([
            row["table_schema"],
            row["table_name"],
            str(col_counts[key]),
            ", ".join(pk_map[key]) or "",
            ", ".join(sorted(set(references[key]))) or "",
            ", ".join(sorted(set(referenced_by[key]))) or "",
        ])
    add_table(document, ["Schema", "Table", "Columns", "Primary Key", "References", "Referenced By"], inventory_rows, widths=[0.85, 1.35, 0.55, 1.0, 1.4, 1.4])

    document.add_page_break()
    document.add_heading("Tables By Schema", level=1)
    for schema in SCHEMA_ORDER:
        schema_tables = [row for row in tables if row["table_schema"] == schema]
        if not schema_tables:
            continue
        document.add_heading(schema, level=2)
        rows = []
        for table in schema_tables:
            key = table_key(table)
            table_columns = [col for col in columns if col["table_schema"] == schema and col["table_name"] == table["table_name"]]
            column_summary = "; ".join(
                f"{col['column_name']} {col['data_type']}{' not null' if col['is_nullable'] == 'NO' else ''}"
                for col in table_columns
            )
            rows.append([
                table["table_name"],
                ", ".join(pk_map[key]) or "",
                str(col_counts[key]),
                column_summary,
            ])
        add_table(document, ["Table", "Primary Key", "Column Count", "Columns"], rows, widths=[1.35, 1.0, 0.75, 3.4])

    document.add_page_break()
    document.add_heading("Foreign-Key Relationships", level=1)
    fk_rows = []
    for fk in foreign_keys:
        fk_rows.append([
            fk["constraint_name"],
            f"{fk['source_schema']}.{fk['source_table']}.{fk['source_column']}",
            f"{fk['target_schema']}.{fk['target_table']}.{fk['target_column']}",
        ])
    add_table(document, ["Constraint", "Source", "Target"], fk_rows, widths=[2.0, 2.2, 2.3])

    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "Beacon database documentation"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(DOCX_PATH)


def main() -> None:
    tables = read_csv("tables.csv")
    columns = read_csv("columns.csv")
    primary_keys = read_csv("primary_keys.csv")
    foreign_keys = read_csv("foreign_keys.csv")
    build_visual(tables, foreign_keys)
    build_docx(tables, columns, primary_keys, foreign_keys)
    print(PNG_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
